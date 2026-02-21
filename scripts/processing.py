#!/usr/bin/env python3
"""
Hybrid preprocessing with agentic chunking:
- CSV rows → one chunk each
- Word/text docs → semantically chunked (agentic)
"""

import os
import json
import pandas as pd
from docx import Document
from langchain_experimental.text_splitter import SemanticChunker
from langchain_ollama import OllamaEmbeddings

RAW_DIR = "data/raw"
OUT_DIR = "data/processed"
os.makedirs(OUT_DIR, exist_ok=True)

# Align with downstream expectation used by vectorstore builder
rags_out = os.path.join(OUT_DIR, "documents_for_rag_final.jsonl")

embeddings = OllamaEmbeddings(model="nomic-embed-text")
# Use a smaller breakpoint threshold to create smaller, more focused chunks
chunker = SemanticChunker(embeddings, breakpoint_threshold_type="percentile", breakpoint_threshold_amount=50)

def clean_text(s: str):
    return " ".join(str(s).split()).strip() if s else ""

def process_faqs():
    path = os.path.join(RAW_DIR, "Faqs.csv")
    if not os.path.exists(path):
        print("[FAQs] Not found")
        return []

    df = pd.read_csv(path, dtype=str)
    cols = {c.lower(): c for c in df.columns}
    qcol, acol = cols.get("question"), cols.get("answer")
    catcol = cols.get("category")

    docs = []
    for i, row in df.iterrows():
        q, a = clean_text(row.get(qcol)), clean_text(row.get(acol))
        cat = clean_text(row.get(catcol, "")) if catcol else ""
        if not (q and a):
            continue
        content = f"Q: {q}\nA: {a}"
        docs.append({
            "id": f"faq-{i}",
            "section": "faq",
            "title": q,
            "content": content,
            "metadata": {"category": cat, "source": "Faqs.csv"}
        })
    print(f"[FAQs] {len(docs)} entries processed")
    return docs

def process_internet():
    path = os.path.join(RAW_DIR, "internet_data.csv")
    if not os.path.exists(path):
        print("[Internet] Not found")
        return []

    df = pd.read_csv(path, dtype=str)
    docs = []
    # Flexible column access with fallbacks
    col = {c.lower(): c for c in df.columns}
    type_col = col.get("internet type", "Internet type")
    bundle_type_col = col.get("internet bundle type", "Internet Bundle Type")
    bundle_col = col.get("internet bundle", "Internet Bundle")
    price_col = col.get("price(egp)", "Price(EGP)")
    dial_col = col.get("to subscribe call", "To subscribe call")
    volume_col = col.get("inclusive volume(mbs)", "Inclusive Volume(MBS)")
    gift_col = col.get("gift", "Gift")
    days_col = col.get("duration_days", "Duration_days")
    hours_col = col.get("duration_hours", "Duration_hours")
    speed_col = col.get("internet speed (mb)", "Internet Speed (MB)")

    for i, row in df.iterrows():
        internet_type = clean_text(row.get(type_col, ""))
        bundle_type = clean_text(row.get(bundle_type_col, ""))
        package = clean_text(row.get(bundle_col, ""))
        price = clean_text(row.get(price_col, ""))
        quota = clean_text(row.get(volume_col, ""))
        validity_days = clean_text(row.get(days_col, ""))
        validity_hours = clean_text(row.get(hours_col, ""))
        dial = clean_text(row.get(dial_col, ""))
        gift = clean_text(row.get(gift_col, ""))
        speed = clean_text(row.get(speed_col, ""))

        # Build rich, queryable content with synonyms/keywords to help retrieval
        lines = []
        if internet_type:
            lines.append(f"Internet type: {internet_type}")
        if bundle_type:
            lines.append(f"Bundle family: {bundle_type}")
        if package:
            lines.append(f"Package name: {package}")
        if price:
            lines.append(f"Price (EGP): {price}")
        if quota:
            lines.append(f"Data quota (MB): {quota}")
        if speed:
            lines.append(f"Speed (Mb): {speed}")
        if validity_days:
            lines.append(f"Validity (days): {validity_days}")
        if validity_hours:
            lines.append(f"Validity (hours): {validity_hours}")
        if gift:
            lines.append(f"Gift/Bonus: {gift}")
        if dial:
            lines.append(f"Subscribe shortcode: {dial}")

        # Add descriptive sentence and keywords to boost semantic match for typical queries
        qualifiers = []
        if internet_type:
            qualifiers.append(internet_type)
        if bundle_type:
            qualifiers.append(bundle_type)

        descriptive = (
            f"This is a mobile internet bundle plan offering {quota or 'a data allowance'} MB data"
            if (internet_type.lower() == "mobile internet" if internet_type else False) else
            f"This is an internet bundle plan offering {quota or 'a data allowance'} MB data"
        )

        keywords = "Keywords: mobile internet, internet bundle, data plan, package, offer, pricing, quota, MB, GB, subscribe, shortcode, #222#, GO, Social, Video, Amazon, Play, @Home, monthly, weekly, daily"

        details = "\n".join(lines + [descriptive, keywords])

        docs.append({
            "id": f"internet-{i}",
            "section": "internet",
            "title": package or (bundle_type or "Internet Package"),
            "content": details,
            "metadata": {
                "source": "internet_data.csv",
                "package_name": package,
                "internet_type": internet_type,
                "bundle_type": bundle_type,
                "dial": dial
            }
        })
    print(f"[Internet] {len(docs)} entries processed")
    return docs

def process_docx():
    path = os.path.join(RAW_DIR, "orange_document.docx")
    if not os.path.exists(path):
        print("[Docx] Not found")
        return []

    doc = Document(path)
    paras = [clean_text(p.text) for p in doc.paragraphs if clean_text(p.text)]
    text = "\n\n".join(paras)

    # semantic (agentic) chunking
    from langchain_core.documents import Document as LCDocument
    doc_obj = LCDocument(page_content=text, metadata={"source": "orange_document.docx"})
    chunks = chunker.split_documents([doc_obj])

    docs = []
    for i, ch in enumerate(chunks):
        # Add rich metadata to help retrieval
        content = ch.page_content
        metadata = dict(ch.metadata)
        
        # Tag documents containing tariff plans with minutes
        if "Minutes to any network" in content or "minutes to Orange network" in content:
            metadata["has_minutes"] = "true"
            metadata["plan_type"] = "tariff_plan"
        
        # Tag specific plan types
        if "PREMIER" in content:
            metadata["plan_category"] = "PREMIER"
        elif "ALO" in content.upper() or "Alo" in content:
            metadata["plan_category"] = "ALO"
        elif "FREEmax" in content:
            metadata["plan_category"] = "FREEmax"
        elif "GO" in content and ("GO " in content or "GO@" in content):
            metadata["plan_category"] = "GO"
        
        docs.append({
            "id": f"doc-{i}",
            "section": "document",
            "title": "orange_document",
            "content": ch.page_content,
            "metadata": metadata
        })
    print(f"[Docx] {len(docs)} semantic chunks")
    return docs

if __name__ == "__main__":
    all_docs = []
    all_docs.extend(process_faqs())
    all_docs.extend(process_internet())
    all_docs.extend(process_docx())

    # save JSONL
    with open(rags_out, "w", encoding="utf-8") as f:
        for d in all_docs:
            f.write(json.dumps(d, ensure_ascii=False) + "\n")

    print(f"[RAG] Wrote {len(all_docs)} docs to {rags_out}")
    print("Preprocessing complete.")
